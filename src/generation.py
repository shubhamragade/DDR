
import io
import markdown
import re
from xhtml2pdf import pisa
from docx import Document

def generate_ddr_markdown(structured_data: dict, template_style: str) -> str:
    """
    Formats the structured data into the final Markdown report.
    """
    return structured_data.get('report_content', "Error: No report content generated.")

def generate_pdf(markdown_content: str) -> bytes:
    """
    Converts Markdown content to PDF bytes using xhtml2pdf.
    Uses CSS for professional layout and alignment.
    """
    processed = re.sub(r'([^\n])(#{1,3}\s)', r'\1\n\n\2', markdown_content)
    
    # 2. Fix merged bullet points (e.g. text- **Area**)
    processed = re.sub(r'([^\n])(-\s\*\*)', r'\1\n\2', processed)
    
    # 3. Remove redundant empty headers like #\n# or # #
    processed = re.sub(r'(^|\n)#\s*\n#', r'\1#', processed)
    processed = re.sub(r'#\s*#', r'#', processed)
    
    # 4. Fix headers merged with previous headers (e.g. # Header## SubHeader)
    processed = re.sub(r'(#[^#\n]+)(#{1,3}\s)', r'\1\n\n\2', processed)

    # Convert Markdown to HTML
    html_content = markdown.markdown(
        processed, 
        extensions=['extra', 'nl2br', 'sane_lists']
    )
    
    # Define CSS for professional alignment and styling
    css = """
    @page {
        size: letter;
        margin: 1in;
    }
    body {
        font-family: Helvetica, Arial, sans-serif;
        font-size: 10pt;
        line-height: 1.5;
        color: #333;
        text-align: justify;
    }
    h1 { font-size: 18pt; margin-top: 20pt; margin-bottom: 10pt; color: #000; text-align: left; }
    h2 { font-size: 14pt; margin-top: 15pt; margin-bottom: 8pt; color: #333; text-align: left; border-bottom: 0.5pt solid #ccc; font-weight: bold; }
    h3 { font-size: 12pt; margin-top: 12pt; margin-bottom: 6pt; color: #444; text-align: left; font-weight: bold; }
    p { margin-bottom: 8pt; }
    ul, ol { margin-bottom: 10pt; margin-left: 20pt; }
    li { margin-bottom: 4pt; }
    hr { border: 0.5pt solid #ccc; margin: 15pt 0; }
    b, strong { color: #000; font-weight: bold; }
    """
    
    # Wrap HTML in a full document structure
    full_html = f"""
    <html>
    <head>
        <style>{css}</style>
    </head>
    <body>
        <div class="content">
            {html_content}
        </div>
    </body>
    </html>
    """
    
    buffer = io.BytesIO()
    try:
        # Create PDF
        pisa_status = pisa.CreatePDF(
            full_html,
            dest=buffer
        )
        
        if pisa_status.err:
            print(f"PDF Generation Error (pisa): {pisa_status.err}")
            return b""
            
        return buffer.getvalue()
    except Exception as e:
        print(f"PDF Generation Exception: {e}")
        return b""

def generate_docx(markdown_content: str) -> bytes:
    """
    Converts Markdown content to a Word Document (bytes).
    Uses a simple custom parser to map Markdown syntax to Word styles.
    """
    doc = Document()
    
    # Basic Markdown parsing
    lines = markdown_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            # Heading 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            # Heading 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Heading 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- ') or line.startswith('* '):
            # List item
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.startswith('1. '):
            # Numbered list (simple check)
            doc.add_paragraph(line[3:], style='List Number')
        elif line.startswith('**') and line.endswith('**'):
            # Bold paragraph
            p = doc.add_paragraph()
            p.add_run(line[2:-2]).bold = True
        else:
            # Regular paragraph
            doc.add_paragraph(line)
            
    # Save to buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    
    return docx_buffer.getvalue()
